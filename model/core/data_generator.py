import os
import random
from pathlib import Path
import pandas as pd
import numpy as np
from faker import Faker
from PyPDF2 import PdfWriter, PdfReader
from io import BytesIO
import docx
from tqdm import tqdm
import argparse
from PIL import Image, ImageDraw, ImageFont
import textwrap
from datetime import datetime, timedelta
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class SyntheticDataGenerator:
    def __init__(self, output_dir="files/synthetic"):
        self.fake = Faker()
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.document_types = {
            "drivers_license": {
                "keywords": ["driver", "license", "licence", "driving licence", "driving license", "driver's license", "driver's licence", "identification", "ID", "operator", "permit", "DOB", "date of birth", "class", "issue date", "expiration", "expires", "restrictions", "endorsements", "organ donor", "DVLA", "DL", "driving", "provisional", "wheeler", "vehicle", "motorist", "number", "license number", "licence number", "state", "sex", "gender", "height", "weight", "eyes", "eye color", "hair", "hair color", "address", "street", "city", "zip", "signature", "hawaii", "honolulu", "peace", "issue", "birth date", "valid", "status", "type"],
                "naming_patterns": ["DL", "drivers_license", "license", "id_card", "identification"]
            },
            "bank_statement": {
                "keywords": ["account", "balance", "transaction", "statement", "deposit", "withdraw", "bank", "checking", "savings", "beginning balance", "ending balance", "ATM", "credit", "debit", "ROUTING", "ACCOUNT NO"],
                "naming_patterns": ["statement", "bank_statement", "account", "transactions", "monthly_statement", "checking_statement", "savings_statement"]
            },
            "invoice": {
                "keywords": ["invoice", "bill", "payment", "due date", "amount due", "total", "subtotal", "tax", "invoice number", "purchase order", "item", "quantity", "unit price", "amount", "terms", "ship to", "bill to"],
                "naming_patterns": ["invoice", "bill", "receipt", "payment", "invoice_no", "inv", "billing"]
            },
            "tax_return": {
                "keywords": ["tax", "return", "IRS", "income", "deduction", "filing", "W-2", "1099", "Form 1040", "exemption", "refund", "tax year", "adjusted gross income", "taxable income", "tax due", "withholding"],
                "naming_patterns": ["tax_return", "taxes", "irs", "form1040", "tax_filing", "1040", "tax_year"]
            },
            "medical_record": {
                "keywords": ["patient", "diagnosis", "prescription", "doctor", "hospital", "medical", "treatment", "health", "insurance", "medication", "allergies", "symptoms", "vital signs", "medical history", "physical examination"],
                "naming_patterns": ["medical", "health_record", "patient", "hospital", "med_record", "clinical", "health_summary"]
            },
            "insurance_claim": {
                "keywords": ["claim", "policy", "insurance", "coverage", "premium", "beneficiary", "policyholder", "insurer", "claim number", "incident", "damage", "loss", "liability", "deductible", "coverage limits"],
                "naming_patterns": ["insurance", "claim", "policy", "claim_form", "insurance_claim", "policy_claim"]
            }
        }
        
        self.supported_formats = ["pdf", "docx", "jpg", "png", "csv"]
        self.background_colors = [(255, 255, 255), (250, 250, 250), (245, 245, 245), (240, 240, 240)]
        self.text_colors = [(0, 0, 0), (50, 50, 50), (25, 25, 112), (0, 0, 128)]
        
        self.banks = ["Chase Bank", "Bank of America", "Wells Fargo", "Citibank", "Capital One", "TD Bank", "PNC Bank", "US Bank"]
        self.companies = ["ABC Corporation", "XYZ Industries", "Acme Supplies", "Tech Solutions Inc.", "Global Services LLC", "Premier Products", "Elite Manufacturing", "Innovative Systems"]
    
    def _generate_drivers_license_content(self):
        state = self.fake.state_abbr()
        dl_number = f"{state}{self.fake.numerify('##-###-####')}"
        issue_date = self.fake.date_between(start_date="-5y", end_date="today")
        expiration_date = issue_date + timedelta(days=365*4 + random.randint(0, 365))
        
        full_name = self.fake.name()
        address = self.fake.street_address()
        city_state_zip = f"{self.fake.city()}, {state} {self.fake.zipcode()}"
        dob = self.fake.date_of_birth(minimum_age=16, maximum_age=90).strftime("%m/%d/%Y")
        
        height = f"{random.randint(4, 6)}'{random.randint(0, 11)}\""
        weight = f"{random.randint(100, 250)} LBS"
        eye_color = random.choice(["BRN", "BLU", "GRN", "HZL", "GRY"])
        hair_color = random.choice(["BRN", "BLK", "BLN", "RED", "GRY"])
        license_class = random.choice(["A", "B", "C", "D"])
        restrictions = random.choice(["NONE", "CORRECTIVE LENSES", "DAYLIGHT DRIVING ONLY"])
        
        content = f"""
{state} DRIVER LICENSE
DL {dl_number}
CLASS {license_class}

NAME: {full_name}
ADDRESS: {address}
         {city_state_zip}
DOB: {dob}
ISSUED: {issue_date.strftime("%m/%d/%Y")}
EXPIRES: {expiration_date.strftime("%m/%d/%Y")}
HT: {height}
WT: {weight}
EYES: {eye_color}
HAIR: {hair_color}
SEX: {random.choice(['M', 'F'])}
RESTRICTIONS: {restrictions}
"""
        return content
    
    def _generate_bank_statement_content(self):
        bank_name = random.choice(self.banks)
        account_number = f"XXXX-XXXX-{self.fake.numerify('####')}"
        routing_number = self.fake.numerify("#########")
        statement_date = self.fake.date_between(start_date="-3m", end_date="today")
        period_start = statement_date.replace(day=1)
        period_end = statement_date
        
        customer_name = self.fake.name()
        address = self.fake.street_address()
        city_state_zip = f"{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}"
        
        beginning_balance = round(random.uniform(500, 10000), 2)
        ending_balance = beginning_balance
        
        transactions = []
        for _ in range(random.randint(5, 15)):
            transaction_date = self.fake.date_between(start_date=period_start, end_date=period_end)
            description = random.choice([
                f"DEPOSIT", 
                f"WITHDRAWAL", 
                f"ATM WITHDRAWAL", 
                f"DIRECT DEPOSIT - {self.fake.company()}", 
                f"POS PURCHASE - {self.fake.company()}", 
                f"ONLINE PAYMENT - {self.fake.company()}", 
                f"CHECK #{self.fake.numerify('####')}", 
                f"TRANSFER TO {self.fake.name()}", 
                f"ACCOUNT FEE"
            ])
            
            if "DEPOSIT" in description or "DIRECT DEPOSIT" in description:
                amount = round(random.uniform(100, 2000), 2)
                ending_balance += amount
            else:
                amount = round(random.uniform(10, 500), 2)
                ending_balance -= amount
            
            transactions.append({
                "date": transaction_date.strftime("%m/%d/%Y"),
                "description": description,
                "amount": amount,
                "balance": round(ending_balance, 2)
            })
        
        content = f"""
{bank_name}
ACCOUNT STATEMENT

CUSTOMER: {customer_name}
{address}
{city_state_zip}

ACCOUNT NUMBER: {account_number}
ROUTING NUMBER: {routing_number}
STATEMENT PERIOD: {period_start.strftime("%m/%d/%Y")} to {period_end.strftime("%m/%d/%Y")}

BEGINNING BALANCE: ${beginning_balance:.2f}
ENDING BALANCE: ${ending_balance:.2f}

TRANSACTION HISTORY:
"""
        
        sorted_transactions = sorted(transactions, key=lambda x: x["date"])
        for t in sorted_transactions:
            content += f"\n{t['date']} {t['description']:<40} ${t['amount']:.2f} ${t['balance']:.2f}"
        
        return content
    
    def _generate_invoice_content(self):
        company = random.choice(self.companies)
        customer = self.fake.company()
        invoice_number = f"INV-{self.fake.numerify('#####')}"
        invoice_date = self.fake.date_between(start_date="-3m", end_date="today")
        due_date = invoice_date + timedelta(days=random.choice([15, 30, 45, 60]))
        
        company_address = f"{self.fake.street_address()}\n{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}"
        customer_address = f"{self.fake.street_address()}\n{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}"
        
        items = []
        subtotal = 0
        for _ in range(random.randint(1, 6)):
            item_name = random.choice([
                "Professional Services", "Consulting", "Product Development", "Support Hours",
                "Design Services", "Website Maintenance", "Software License", "Hardware Purchase",
                "Training Session", "Project Management", "Monthly Subscription", "Data Storage"
            ])
            quantity = random.randint(1, 10)
            unit_price = round(random.uniform(50, 500), 2)
            amount = round(quantity * unit_price, 2)
            subtotal += amount
            
            items.append({
                "item": item_name,
                "quantity": quantity,
                "unit_price": unit_price,
                "amount": amount
            })
            
        tax_rate = random.uniform(0.05, 0.095)
        tax = round(subtotal * tax_rate, 2)
        total = subtotal + tax
        
        content = f"""
{company}
{company_address}
Phone: {self.fake.phone_number()}
Email: {self.fake.email()}

INVOICE

INVOICE #: {invoice_number}
DATE: {invoice_date.strftime("%m/%d/%Y")}
DUE DATE: {due_date.strftime("%m/%d/%Y")}

BILL TO:
{customer}
{customer_address}

DESCRIPTION                      QUANTITY    UNIT PRICE     AMOUNT
"""
        
        for item in items:
            content += f"\n{item['item']:<30} {item['quantity']:<10} ${item['unit_price']:<14.2f} ${item['amount']:.2f}"
            
        content += f"""

                                             SUBTOTAL:     ${subtotal:.2f}
                                             TAX ({tax_rate*100:.1f}%):     ${tax:.2f}
                                             TOTAL:        ${total:.2f}

PAYMENT TERMS: Net {(due_date - invoice_date).days} days
PAYMENT METHODS: Check, Credit Card, Bank Transfer

Thank you for your business!
"""
        return content
    
    def _generate_tax_return_content(self):
        tax_year = random.randint(datetime.now().year - 3, datetime.now().year - 1)
        taxpayer_name = self.fake.name()
        ssn = f"XXX-XX-{self.fake.numerify('####')}"
        address = self.fake.street_address()
        city_state_zip = f"{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}"
        filing_status = random.choice(["Single", "Married Filing Jointly", "Married Filing Separately", "Head of Household"])
        
        wages = round(random.uniform(25000, 120000), 2)
        interest_income = round(random.uniform(0, 5000), 2)
        dividend_income = round(random.uniform(0, 5000), 2)
        business_income = round(random.uniform(0, 30000), 2)
        capital_gains = round(random.uniform(0, 10000), 2)
        
        total_income = wages + interest_income + dividend_income + business_income + capital_gains
        adjustments = round(random.uniform(0, 5000), 2)
        adjusted_gross_income = total_income - adjustments
        
        standard_deduction = 12000 if filing_status == "Single" else 24000
        taxable_income = max(0, adjusted_gross_income - standard_deduction)
        
        tax_rate = 0.22 if taxable_income > 40000 else 0.12
        tax_due = round(taxable_income * tax_rate, 2)
        
        tax_withheld = round(wages * 0.18, 2)
        tax_owed = round(tax_due - tax_withheld, 2)
        
        content = f"""
FORM 1040: U.S. INDIVIDUAL INCOME TAX RETURN

TAX YEAR: {tax_year}

TAXPAYER INFORMATION:
  Name: {taxpayer_name}
  SSN: {ssn}
  Address: {address}
  {city_state_zip}
  Filing Status: {filing_status}

INCOME:
  Wages, salaries, tips, etc.: ${wages:.2f}
  Interest income: ${interest_income:.2f}
  Dividend income: ${dividend_income:.2f}
  Business income: ${business_income:.2f}
  Capital gains: ${capital_gains:.2f}
  TOTAL INCOME: ${total_income:.2f}

ADJUSTMENTS TO INCOME:
  Total adjustments: ${adjustments:.2f}
  ADJUSTED GROSS INCOME: ${adjusted_gross_income:.2f}

DEDUCTIONS:
  Standard deduction: ${standard_deduction:.2f}
  TAXABLE INCOME: ${taxable_income:.2f}

TAX AND CREDITS:
  Tax (rate: {tax_rate*100:.1f}%): ${tax_due:.2f}
  Total tax: ${tax_due:.2f}

PAYMENTS:
  Federal income tax withheld: ${tax_withheld:.2f}
  TOTAL PAYMENTS: ${tax_withheld:.2f}

REFUND OR AMOUNT YOU OWE:
  TAX OWED: ${max(0, tax_owed):.2f}
  REFUND AMOUNT: ${max(0, -tax_owed):.2f}

I declare under penalty of perjury that this return is true, correct, and complete.

Signature: ____________________ Date: _____________
"""
        return content
    
    def _generate_medical_record_content(self):
        patient_name = self.fake.name()
        dob = self.fake.date_of_birth(minimum_age=18, maximum_age=90).strftime("%m/%d/%Y")
        mrn = f"MRN-{self.fake.numerify('#######')}"
        insurance = random.choice(["Blue Cross Blue Shield", "Aetna", "UnitedHealthcare", "Cigna", "Humana", "Medicare"])
        policy_number = f"POL-{self.fake.numerify('########')}"
        
        visit_date = self.fake.date_between(start_date="-1y", end_date="today")
        provider = f"Dr. {self.fake.name()}"
        facility = f"{self.fake.company()} Medical Center"
        
        symptoms = random.sample([
            "Fever", "Headache", "Cough", "Fatigue", "Nausea", "Abdominal pain", 
            "Chest pain", "Back pain", "Dizziness", "Shortness of breath"
        ], random.randint(1, 3))
        
        vitals = {
            "Temperature": f"{round(random.uniform(97.5, 99.8), 1)}°F",
            "Blood Pressure": f"{random.randint(110, 140)}/{random.randint(60, 90)} mmHg",
            "Heart Rate": f"{random.randint(60, 100)} bpm",
            "Respiratory Rate": f"{random.randint(12, 18)} breaths/min",
            "Oxygen Saturation": f"{random.randint(95, 100)}%"
        }
        
        conditions = random.sample([
            "Hypertension", "Type 2 Diabetes", "Asthma", "Coronary Artery Disease", 
            "Hypothyroidism", "Osteoarthritis", "GERD", "Chronic Sinusitis"
        ], random.randint(0, 2))
        
        medications = random.sample([
            "Lisinopril 10mg daily", "Metformin 500mg twice daily", "Albuterol inhaler as needed",
            "Atorvastatin 20mg daily", "Levothyroxine 50mcg daily", "Omeprazole 20mg daily"
        ], random.randint(0, 3))
        
        allergies = random.sample(["Penicillin", "Sulfa drugs", "NSAIDs", "Shellfish", "No known allergies"], 1)
        
        diagnoses = random.sample([
            "Acute Bronchitis", "Hypertension", "Gastroenteritis", "Migraine", 
            "Urinary Tract Infection", "Anxiety", "Lumbar Strain", "Pharyngitis"
        ], random.randint(1, 2))
        
        treatment = random.sample([
            "Prescribed antibiotics", "Rest and hydration", "Physical therapy", 
            "Medication adjustment", "Follow-up in 2 weeks", "Specialist referral"
        ], random.randint(1, 3))
        
        treatment_text = "\n".join(treatment)
        medications_text = "\n".join(medications) if medications else "None"
        conditions_text = ", ".join(conditions) if conditions else "None"
        
        content = f"""
MEDICAL RECORD

PATIENT INFORMATION:
Name: {patient_name}
DOB: {dob}
MRN: {mrn}
Insurance: {insurance}
Policy Number: {policy_number}

VISIT INFORMATION:
Date of Visit: {visit_date.strftime("%m/%d/%Y")}
Provider: {provider}
Facility: {facility}

CHIEF COMPLAINT:
{", ".join(symptoms)}

VITAL SIGNS:
"""
        
        for vital, value in vitals.items():
            content += f"{vital}: {value}\n"
            
        content += f"""
MEDICAL HISTORY:
{conditions_text}

CURRENT MEDICATIONS:
{medications_text}

ALLERGIES:
{allergies[0]}

ASSESSMENT:
{", ".join(diagnoses)}

TREATMENT PLAN:
{treatment_text}

PROVIDER SIGNATURE: ____________________
"""
        return content
    
    def _generate_insurance_claim_content(self):
        insurance_company = random.choice(["State Farm", "Allstate", "Geico", "Progressive", "Farmers", "Liberty Mutual"])
        claim_number = f"CLAIM-{self.fake.numerify('#######')}"
        policy_number = f"POL-{self.fake.numerify('########')}"
        date_of_loss = self.fake.date_between(start_date="-6m", end_date="-1d")
        date_reported = date_of_loss + timedelta(days=random.randint(0, 10))
        
        policyholder = self.fake.name()
        address = self.fake.street_address()
        city_state_zip = f"{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}"
        phone = self.fake.phone_number()
        
        claim_types = {
            "Auto": ["Collision", "Comprehensive", "Liability"],
            "Home": ["Property Damage", "Theft", "Liability", "Water Damage", "Fire Damage"],
            "Health": ["Emergency Care", "Surgery", "Specialist Visit", "Diagnostic Tests"],
            "Life": ["Death Benefit"]
        }
        
        claim_type = random.choice(list(claim_types.keys()))
        incident_type = random.choice(claim_types[claim_type])
        
        if claim_type == "Auto":
            vehicle = f"{random.randint(2010, 2023)} {self.fake.company()} {random.choice(['Sedan', 'SUV', 'Truck', 'Crossover'])}"
            vin = self.fake.lexify(text="?" * 17, letters="ABCDEFGHJKLMNPRSTUVWXYZ0123456789")
            incident_location = f"{self.fake.street_address()}, {self.fake.city()}, {self.fake.state_abbr()}"
            description = random.choice([
                f"Vehicle was involved in a collision with another vehicle at an intersection.",
                f"Vehicle was parked and hit by an unknown driver.",
                f"Vehicle was damaged by hail during a storm.",
                f"Vehicle was stolen from a parking lot."
            ])
        elif claim_type == "Home":
            property_address = f"{address}, {city_state_zip}"
            description = random.choice([
                f"Water damage from broken pipe in the kitchen.",
                f"Theft of personal property from locked residence.",
                f"Wind damage to roof during storm.",
                f"Fire damage to portion of the home.",
                f"Damage from fallen tree."
            ])
        else:
            description = f"Claim related to {incident_type} services provided on {date_of_loss.strftime('%m/%d/%Y')}."
        
        estimated_damage = round(random.uniform(1000, 20000), 2)
        deductible = round(random.uniform(250, 1000), 2)
        claim_amount = round(estimated_damage - deductible, 2)
        
        content = f"""
INSURANCE CLAIM FORM

{insurance_company}

CLAIM INFORMATION:
Claim Number: {claim_number}
Policy Number: {policy_number}
Date of Loss: {date_of_loss.strftime("%m/%d/%Y")}
Date Reported: {date_reported.strftime("%m/%d/%Y")}
Claim Type: {claim_type}
Incident Type: {incident_type}

POLICYHOLDER INFORMATION:
Name: {policyholder}
Address: {address}
        {city_state_zip}
Phone: {phone}

"""
        
        if claim_type == "Auto":
            content += f"""VEHICLE INFORMATION:
Vehicle: {vehicle}
VIN: {vin}
Incident Location: {incident_location}
"""
        elif claim_type == "Home":
            content += f"""PROPERTY INFORMATION:
Property Address: {property_address}
"""
        
        content += f"""
INCIDENT DESCRIPTION:
{description}

CLAIM DETAILS:
Estimated Damage: ${estimated_damage:.2f}
Deductible: ${deductible:.2f}
Claim Amount: ${claim_amount:.2f}

CLAIM STATUS: {random.choice(["Pending", "Under Review", "Approved", "Paid"])}

ADJUSTER: {self.fake.name()}
NOTES: {self.fake.sentence()}

POLICYHOLDER SIGNATURE: __________________ DATE: __________
"""
        return content
    
    def _generate_random_text(self, doc_type):
        content_generators = {
            "drivers_license": self._generate_drivers_license_content,
            "bank_statement": self._generate_bank_statement_content,
            "invoice": self._generate_invoice_content,
            "tax_return": self._generate_tax_return_content,
            "medical_record": self._generate_medical_record_content,
            "insurance_claim": self._generate_insurance_claim_content
        }
        
        if doc_type in content_generators:
            content = content_generators[doc_type]()
        else:
            content = self.fake.text(max_nb_chars=2000)
            
        keywords = self.document_types[doc_type]["keywords"]
        random_words = self.fake.words(nb=10)
        
        injected_words = keywords * 3 + random_words
        
        for _ in range(random.randint(20, 40)):
            word = random.choice(injected_words)
            position = random.randint(0, max(0, len(content) - len(word) - 1))
            content = content[:position] + " " + word + " " + content[position:]
            
        return content
    
    def _get_random_filename(self, doc_type, file_format):
        patterns = self.document_types[doc_type]["naming_patterns"]
        pattern = random.choice(patterns)
        
        if random.random() < 0.3:
            filename = f"{pattern}_{self.fake.random_number(digits=4)}.{file_format}"
        elif random.random() < 0.6:
            filename = f"{self.fake.word()}_{pattern}.{file_format}"
        else:
            filename = f"{self.fake.random_number(digits=2)}_{pattern}_{self.fake.date()}.{file_format}"
            
        if random.random() < 0.2:
            filename = filename.upper()
        elif random.random() < 0.4:
            filename = filename.replace("_", "")
            
        return filename
    
    def _generate_pdf(self, doc_type, poorly_named=False):
        content = self._generate_random_text(doc_type)
        
        if poorly_named:
            filename = f"{self.fake.word()}_{self.fake.random_number(digits=3)}.pdf"
        else:
            filename = self._get_random_filename(doc_type, "pdf")
            
        file_path = self.output_dir / filename
        
        c = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter
        
        title = doc_type.replace("_", " ").title()
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, height - 72, title)
        
        y_position = height - 100
        c.setFont("Helvetica", 10)
        
        for line in content.split('\n'):
            if not line.strip():
                y_position -= 12
                continue
                
            wrapped_lines = textwrap.wrap(line, width=100)
            for wrapped_line in wrapped_lines:
                c.drawString(72, y_position, wrapped_line)
                y_position -= 12
                
                if y_position < 72:
                    c.showPage()
                    y_position = height - 72
                    c.setFont("Helvetica", 10)
        
        c.save()
        
        return {
            "filename": filename,
            "path": str(file_path),
            "content": content,
            "type": doc_type,
            "poorly_named": poorly_named
        }
    
    def _generate_docx(self, doc_type, poorly_named=False):
        document = docx.Document()
        content = self._generate_random_text(doc_type)
        
        document.add_heading(doc_type.replace("_", " ").title(), 0)
        document.add_paragraph(content)
        
        if poorly_named:
            filename = f"{self.fake.word()}_{self.fake.random_number(digits=3)}.docx"
        else:
            filename = self._get_random_filename(doc_type, "docx")
            
        file_path = self.output_dir / filename
        document.save(str(file_path))
        
        return {
            "filename": filename,
            "path": str(file_path),
            "content": content,
            "type": doc_type,
            "poorly_named": poorly_named
        }
        
    def _generate_image(self, doc_type, image_format="jpg", poorly_named=False):
        width, height = 1200, 1550
        
        content = self._generate_random_text(doc_type)
        
        background_color = random.choice(self.background_colors)
        text_color = random.choice(self.text_colors)
        
        image = Image.new('RGB', (width, height), color=background_color)
        draw = ImageDraw.Draw(image)
        
        try:
            font = ImageFont.truetype("arial.ttf", 24)
            title_font = ImageFont.truetype("arial.ttf", 36)
        except:
            font = ImageFont.load_default()
            title_font = ImageFont.load_default()
        
        title = doc_type.replace("_", " ").title()
        draw.text((width//2, 100), title, fill=text_color, font=title_font, anchor="mm")
        
        wrapped_text = textwrap.fill(content, width=65)
        
        y_position = 180
        for line in wrapped_text.split('\n'):
            draw.text((100, y_position), line, fill=text_color, font=font)
            y_position += 30
        
        if poorly_named:
            filename = f"{self.fake.word()}_{self.fake.random_number(digits=3)}.{image_format}"
        else:
            filename = self._get_random_filename(doc_type, image_format)
            
        file_path = self.output_dir / filename
        image.save(str(file_path))
        
        return {
            "filename": filename,
            "path": str(file_path),
            "content": content,
            "type": doc_type,
            "poorly_named": poorly_named
        }
    
    def _generate_csv(self, doc_type, poorly_named=False):
        content = self._generate_random_text(doc_type)
        
        if poorly_named:
            filename = f"{self.fake.word()}_{self.fake.random_number(digits=3)}.csv"
        else:
            filename = self._get_random_filename(doc_type, "csv")
            
        file_path = self.output_dir / filename
        
        headers = ["id", "date", "type", "description", "value", "keywords"]
        rows = []
        
        for i in range(random.randint(10, 20)):
            keywords = random.sample(self.document_types[doc_type]["keywords"] * 2, 
                                   k=min(random.randint(4, 8), len(self.document_types[doc_type]["keywords"] * 2)))
            keywords_str = "|".join(keywords)
            
            desc_prefix = random.choice([
                f"{doc_type.replace('_', ' ').title()} - ",
                f"Document type: {doc_type.replace('_', ' ')} - ",
                f"{random.choice(self.document_types[doc_type]['keywords']).upper()}: "
            ])
            
            row = {
                "id": f"{self.fake.bothify(text='???###')}",
                "date": self.fake.date_between(start_date="-1y", end_date="today").strftime("%Y-%m-%d"),
                "type": doc_type.replace("_", " ").title(),
                "description": desc_prefix + self.fake.sentence(),
                "value": f"${self.fake.random_number(digits=4)}.{self.fake.random_number(digits=2)}",
                "keywords": keywords_str
            }
            rows.append(row)
        
        with open(file_path, 'w', newline='') as f:
            f.write(','.join(headers) + '\n')
            for row in rows:
                values = []
                for header in headers:
                    value = str(row[header])
                    if ',' in value:
                        value = f'"{value}"'
                    values.append(value)
                f.write(','.join(values) + '\n')
        
        return {
            "filename": filename,
            "path": str(file_path),
            "content": content,
            "type": doc_type,
            "poorly_named": poorly_named
        }
    
    def _generate_file(self, doc_type, file_format, poorly_named=False):
        generators = {
            "pdf": self._generate_pdf,
            "docx": self._generate_docx,
            "jpg": lambda dt, pn: self._generate_image(dt, "jpg", pn),
            "png": lambda dt, pn: self._generate_image(dt, "png", pn),
            "csv": self._generate_csv
        }
        
        if file_format in generators:
            return generators[file_format](doc_type, poorly_named)
        else:
            return None
    
    def generate_dataset(self, num_samples=1000, poorly_named_ratio=0.3):
        data = []
        
        for _ in tqdm(range(num_samples), desc="Generating synthetic files"):
            doc_type = random.choice(list(self.document_types.keys()))
            file_format = random.choice(self.supported_formats)
            
            poorly_named = random.random() < poorly_named_ratio
            
            sample = self._generate_file(doc_type, file_format, poorly_named)
            if sample:
                training_sample = {
                    "path": sample["path"],
                    "content": sample["content"],
                    "type": sample["type"]
                }
                data.append(training_sample)
            
        metadata_df = pd.DataFrame(data)
        metadata_df.to_csv(self.output_dir / "metadata.csv", index=False)
        
        print(f"Generated {len(data)} synthetic files in {self.output_dir}")
        print(f"File type distribution:")
        
        format_counts = {}
        for _, row in metadata_df.iterrows():
            ext = os.path.splitext(row['path'])[1][1:]
            format_counts[ext] = format_counts.get(ext, 0) + 1
            
        for fmt, count in format_counts.items():
            print(f"  {fmt}: {count} files")
            
        return metadata_df

def main():
    parser = argparse.ArgumentParser(description="Generate synthetic data for file classifier training")
    parser.add_argument("--output-dir", type=str, default="files/synthetic",
                      help="Directory to save synthetic data (default: files/synthetic)")
    parser.add_argument("--num-samples", type=int, default=1000,
                      help="Number of synthetic samples to generate (default: 1000)")
    parser.add_argument("--poorly-named-ratio", type=float, default=0.3,
                      help="Ratio of poorly named files (default: 0.3)")
    
    args = parser.parse_args()
    
    generator = SyntheticDataGenerator(output_dir=args.output_dir)
    dataset = generator.generate_dataset(
        num_samples=args.num_samples,
        poorly_named_ratio=args.poorly_named_ratio
    )
    
    print(f"Dataset distribution:\n{dataset['type'].value_counts()}")

if __name__ == "__main__":
    main()